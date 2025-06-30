import streamlit as st
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import numpy as np
import io
import pandas as pd 
from docx.table import Table
import re

# === Funções de análise ===

def analisar_paragrafos(paragraphs, idx_table):
    count_conf = 0
    count_nao_conf = 0
    descricoes = []

    for paragraph in paragraphs:
        texto = paragraph.text
        count_nao_conf += len(re.findall(r"não\s*conforme", texto.lower()))

        # ✔️ Contar "Conforme" quando aparecer após emoji em runs separados
        runs = paragraph.runs
        for i in range(len(runs)):
            texto_emoji = runs[i].text.strip()
            if texto_emoji in ["✔", "✔️", "✓", "✅"]:
                for j in range(i + 1, min(i + 4, len(runs))):
                    if runs[j].text.strip().lower() == "conforme":
                        count_conf += 1
                        break
                break

        # ✔️ Contar "Conforme" quando estiver na mesma run do emoji
        for run in runs:
            texto_run = run.text.strip().lower()
            if any(e in texto_run for e in ["✔", "✔️", "✓", "✅"]) and "conforme" in texto_run:
                count_conf += 1
                break

        # 🟥 Coletar descrições em vermelho
        if "descrição" in texto.lower():
            passou_por_descricao = False
            texto_runs = []

            for run in runs:
                texto_run = run.text
                if "Descrição" in texto_run:
                    passou_por_descricao = True
                elif passou_por_descricao:
                    cor = run.font.color
                    if cor and cor.rgb in [RGBColor(255, 0, 0), RGBColor(238, 0, 0)] and texto_run.strip():
                        texto_runs.append(texto_run)

            if texto_runs:
                descricao_limpinha = texto_runs[0].strip() + ''.join(r.lstrip() for r in texto_runs[1:])
                descricao_limpinha = descricao_limpinha.strip()
                descricoes.append((descricao_limpinha, idx_table))

    return count_conf, count_nao_conf, descricoes

def analisar_tabela(table, idx_table):
    total_conf = 0
    total_nao_conf = 0
    descricoes_encontradas = []

    for row in table.rows:
        for cell in row.cells:
            c_conf, c_nao_conf, descs = analisar_paragrafos(cell.paragraphs, idx_table)
            total_conf += c_conf
            total_nao_conf += c_nao_conf
            descricoes_encontradas.extend(descs)

            for tbl_el in cell._element.xpath(".//w:tbl"):
                try:
                    subtable = Table(tbl_el, cell)
                    c_conf, c_nao_conf, sub_descs = analisar_tabela(subtable, idx_table)
                    total_conf += c_conf
                    total_nao_conf += c_nao_conf
                    descricoes_encontradas.extend(sub_descs)
                except Exception:
                    continue

    return total_conf, total_nao_conf, descricoes_encontradas

# === Interface Streamlit ===

st.title("📄 Analisador de Conformidades em Documento Word")

st.info(
    "🔒 **Aviso de privacidade**:\n\n"
    "Este aplicativo não armazena permanentemente os arquivos enviados. "
    "Todos os documentos são processados apenas durante a sessão atual e são descartados ao final."
)

st.subheader("🔒 Acesso Restrito")
senha_correta = st.secrets["senha"]
senha_digitada = st.text_input("Digite a senha para continuar:", type="password")
if senha_digitada != senha_correta:
    st.warning("Acesso negado. Insira a senha correta.")
    st.stop()

uploaded_file = st.file_uploader("📤 Envie o arquivo Word (.docx)", type="docx")

if uploaded_file:
    st.success("📁 Arquivo carregado com sucesso!")
    doc = Document(uploaded_file)

    count_conforme = 0
    count_nao_conforme = 0
    descricoes_docx = []

    for idx_table, table in enumerate(doc.tables, start=1):
        c_conf, c_nao_conf, descs = analisar_tabela(table, idx_table)
        count_conforme += c_conf
        count_nao_conforme += c_nao_conf
        descricoes_docx.extend(descs)

    st.write(f"✔️ Total 'Conforme': {count_conforme}")
    st.write(f"❌ Total 'Não Conforme': {count_nao_conforme}")

    st.subheader("📝 Descrições Encontradas")

    if descricoes_docx:
        df_descricoes = pd.DataFrame(descricoes_docx, columns=["Descrição", "Figura"])
        st.table(df_descricoes)
    else:
        st.info("Nenhuma descrição em vermelho foi encontrada.")

    # Gerar .xlsx com coluna "Situação"
    df_completo = df_descricoes.copy()
    df_completo["Situação"] = ""

    excel_buffer = io.BytesIO()
    with pd.ExcelWriter(excel_buffer, engine='xlsxwriter') as writer:
        df_completo.to_excel(writer, index=False, sheet_name='Análise')
    excel_buffer.seek(0)

    st.download_button(
        label="📥 Baixar Tabela Completa (.xlsx)",
        data=excel_buffer,
        file_name=uploaded_file.name.replace(".docx", " - Análise.xlsx"),
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

    # Gráfico
    fig, ax = plt.subplots(figsize=(6, 3), subplot_kw=dict(aspect="equal"))
    labels = ["Conforme", "Não conforme"]
    data = [count_conforme, count_nao_conforme]
    colors = ['#4CAF50', '#F44336']

    def func(pct, allvals):
        absolute = int(np.round(pct / 100. * np.sum(allvals)))
        return f"{pct:.1f}%\n({absolute:d})"

    wedges, _, autotexts = ax.pie(
        data, autopct=lambda pct: func(pct, data),
        textprops=dict(color="w"),
        colors=colors
    )
    ax.legend(wedges, labels, title="Situação", loc="center left", bbox_to_anchor=(1, 0, 0.5, 1))
    plt.setp(autotexts, size=8, weight="bold")
    ax.set_title("Análise de Conformidades")
    grafico_path = "grafico_pizza.png"
    plt.savefig(grafico_path)
    st.subheader("📊 Gráfico de Conformidades")
    st.pyplot(fig)
    plt.close()

    # Inserir tabela (2 colunas) no Word
    doc.add_page_break()
    tabela = doc.add_table(rows=len(descricoes_docx) + 1, cols=2)
    try:
        tabela.style = 'Table Grid'
    except KeyError:
        pass
    cabecalhos = ["Descrição", "Figura"]

    for col, texto in enumerate(cabecalhos):
        cell = tabela.cell(0, col)
        cell.text = texto
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(10)
            run.font.bold = True

    for i, (descricao, num_tabela) in enumerate(descricoes_docx, start=1):
        cell_desc = tabela.cell(i, 0)
        run = cell_desc.paragraphs[0].add_run(descricao)
        run.font.size = Pt(10)
        cell_fig = tabela.cell(i, 1)
        cell_fig.text = str(num_tabela)
        for run in cell_fig.paragraphs[0].runs:
            run.font.size = Pt(10)

    # Inserir gráfico
    paragrafo_img = doc.add_paragraph()
    run = paragrafo_img.add_run()
    run.add_picture(grafico_path, width=Inches(5))
    paragrafo_img.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Salvar Word
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.success("✅ Documento atualizado com gráfico e tabela ao final.")
    st.download_button(
        label="📥 Baixar novo Word",
        data=buffer,
        file_name=uploaded_file.name.replace(".docx", " - Análise.docx"),
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
